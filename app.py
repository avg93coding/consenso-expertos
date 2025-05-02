import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import uuid, qrcode, io, hashlib, datetime, base64, copy, os
from scipy import stats
from streamlit_autorefresh import st_autorefresh
import requests
from io import BytesIO
# Reemplaza tus líneas de import de docx por esto:
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


# 1) set_page_config debe ir primero
st.set_page_config(
    page_title="ODDS Epidemiology – Dashboard Consenso de expertos",
    page_icon="https://www.oddsepidemiology.com/favicon.ico",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ——————————————————————————————
# Definición de estilos CSS
def inject_css():
    css = """
    <style>
      .stApp {
        background-color: #F7F7F7 !important;
        color: #333333;
        font-family: 'Segoe UI', Tahoma, Verdana, sans-serif;
      }
      .app-header {
        background-color: #662D91;
        padding: 1.5rem;
        border-radius: 0 0 10px 10px;
        text-align: center;
        color: white;
        margin-bottom: 20px;
      }
      .odds-logo {
        font-size: 2rem;
        font-weight: bold;
        letter-spacing: 1px;
        padding-bottom: 5px;
        border-bottom: 2px solid #F1592A;
        display: inline-block;
      }
      .metric-card {
        width: 140px;
        padding: 12px;
        background: linear-gradient(to bottom right, #662D91, #F1592A);
        color: white;
        border-radius: 8px;
        box-sizing: border-box;
        text-align: center;
      }
      .metric-label {
        font-size: 0.9rem;
        opacity: 0.8;
      }
      .metric-value {
        font-size: 1.4rem;
        font-weight: bold;
        margin-top: 4px;
      }
      .stButton > button {
        background-color: #662D91;
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 5px;
      }
      .stButton > button:hover {
        background-color: #F1592A;
      }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def inject_grid_css():
    grid_css = """
    <style>
      .metric-grid {
        display: grid;
        grid-template-columns: repeat(2, 1fr);
        gap: 12px;
      }
    </style>
    """
    st.markdown(grid_css, unsafe_allow_html=True)

# ——————————————————————————————
# Llamada única a los estilos justo después de los imports
inject_css()
inject_grid_css()

# ——————————————————————————————
# Función auxiliar para generar cada tarjeta
def card_html(label, value):
    return f"""
    <div class="metric-card">
      <div class="metric-label">{label}</div>
      <div class="metric-value">{value}</div>
    </div>
    """


# 3) odds_header(), para mostrar logo y título
def odds_header():
    header_html = """
    <div class="app-header">
      <div class="odds-logo">ODDS EPIDEMIOLOGY</div>
      <div class="odds-subtitle">Sistema de Votación</div>
    </div>
    """
    st.markdown(header_html, unsafe_allow_html=True)

DOMINIOS_GRADE = {
    "prioridad_problema": [
        "No", "Probablemente no", "Probablemente sí", "Sí", "Varía", "No sabemos"
    ],
    "efectos_deseables": [
        "No importante", "Pequeña", "Moderada", "Grande", "Varía", "No se sabe"
    ],
    "efectos_indeseables": [
        "No importante", "Pequeña", "Moderada", "Grande", "Varía", "No se sabe"
    ],
    "certeza_evidencia": [
        "Muy baja", "Baja", "Moderada", "Alta", "No hay estudios incluidos"
    ],
    "balance_efectos": [
        "Favorece al comparador",
        "Probablemente favorece al comparador",
        "No favorece ni al comparador ni a la intervención",
        "Probablemente favorece a la intervención",
        "Favorece la intervención",
        "Es variable",
        "No es posible saber"
    ],
    "recursos": [
        "Costos altos/recursos",
        "Costos moderados/recursos",
        "Costos o ahorro mínimo/recursos insignificantes",
        "Ahorro moderado",
        "Gran ahorro",
        "Variable",
        "No se sabe"
    ],
    "aceptabilidad": [
        "No", "Probablemente no", "Probablemente sí", "Sí", "Varía", "No se sabe"
    ],
    "factibilidad": [
        "No", "Probablemente no", "Probablemente sí", "Sí", "Varía", "No se sabe"
    ],
    "equidad": [
        "Reducido", "Probablemente reducido", "Probablemente no impacta",
        "Probablemente incrementa", "Incrementa", "Varía", "No se sabe"
    ],
}

PREGUNTAS_GRADE = {
    "prioridad_problema":   "¿Constituye el problema una prioridad?",
    "efectos_deseables":    "¿Cuál es la magnitud de los efectos deseados que se prevén?",
    "efectos_indeseables":  "¿Cuál es la magnitud de los efectos no deseados que se prevén?",
    "certeza_evidencia":    "¿Cuál es la certeza global de la evidencia de los efectos?",
    "balance_efectos":      "¿Qué balance entre efectos deseables y no deseables favorece?",
    "recursos":             "¿Cuál es la magnitud de los recursos necesarios (costos)?",
    "aceptabilidad":        "¿Es aceptable la intervención para los grupos clave?",
    "factibilidad":         "¿Es factible la implementación de la intervención?",
    "equidad":              "¿Cuál sería el impacto sobre la equidad en salud?",
}
# ------------------------------------------------------------

import io
import pandas as pd
from scipy import stats

import io
import pandas as pd
import numpy as np
from scipy import stats

def crear_excel_consolidado(store: dict, history: dict) -> io.BytesIO:
    """
    Genera un Excel con tres hojas:
      1) Recomendaciones estándar
      2) Paquetes GRADE
      3) Métricas consolidadas (n, media, mediana, desv. std, % consenso, quórum, estado)
    """
    # — Hoja 1: Recomendaciones estándar —
    filas_std = []
    for code, s in store.items():
        if s.get("tipo", "STD") == "STD":
            for pid, name, vote, com in zip(s["ids"], s["names"], s["votes"], s["comments"]):
                filas_std.append({
                    "Código": code,
                    "Descripción": s["desc"],
                    "Ronda": s["round"],
                    "Creada": s["created_at"],
                    "ID participante": pid,
                    "Nombre": name,
                    "Voto": vote,
                    "Comentario": com
                })
    df_std = pd.DataFrame(filas_std)

    # — Hoja 2: Paquetes GRADE —
    filas_grade = []
    for code, s in store.items():
        if s.get("tipo") == "GRADE_PKG":
            for dom, meta in s["dominios"].items():
                for pid, name, vote, com in zip(meta["ids"], meta["names"], meta["votes"], meta["comments"]):
                    filas_grade.append({
                        "Paquete": code,
                        "Dominio": dom,
                        "ID participante": pid,
                        "Nombre": name,
                        "Voto": vote,
                        "Comentario": com,
                        "Creada": s["created_at"]
                    })
    df_grade = pd.DataFrame(filas_grade)

    # — Hoja 3: Métricas consolidadas —
    filas_metrics = []
    for code, s in store.items():
        votos = [v for v in s["votes"] if isinstance(v, (int, float))]
        n = len(votos)
        media   = np.mean(votos)            if n else np.nan
        std     = np.std(votos, ddof=1)     if n > 1 else 0.0
        mediana = np.median(votos)          if n else np.nan

        # Bootstrap para IC95% de la mediana
        if n >= 2:
            res = stats.bootstrap((votos,), np.median,
                                  confidence_level=0.95,
                                  n_resamples=500,
                                  method="basic")
            lo, hi = res.confidence_interval
        else:
            lo = hi = mediana

        pct_consenso = consensus_pct(votos) * 100
        quorum = s.get("n_participantes", 0)//2 + 1

        # Estado de consenso
        if n < quorum:
            estado = "⚠️ Quórum no alcanzado"
        elif pct_consenso >= 80 and lo >= 7:
            estado = "✅ Consenso alcanzado"
        else:
            estado = "❌ No alcanzó consenso"

        filas_metrics.append({
            "Código":         code,
            "Descripción":    s["desc"],
            "Ronda":          s["round"],
            "Creada":         s["created_at"],
            "Votos totales":  n,
            "Media":          media,
            "Desv. std.":     std,
            "Mediana":        mediana,
            "IC95% (lo)":     lo,
            "IC95% (hi)":     hi,
            "% Consenso":     pct_consenso,
            "Quórum":         quorum,
            "Estado":         estado
        })
    df_metrics = pd.DataFrame(filas_metrics)

    # — Escribir a Excel en memoria —
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_std.to_excel(writer, sheet_name="Recomendaciones", index=False)
        df_grade.to_excel(writer, sheet_name="Paquetes_GRADE", index=False)
        df_metrics.to_excel(writer, sheet_name="Métricas", index=False)
    buffer.seek(0)
    return buffer


def shade_cell(cell, fill_hex: str):
    """
    Aplica un fondo de color (hex sin ‘#’) a una celda de python-docx.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


# Define tus colores corporativos al inicio del fichero
PRIMARY = "#662D91"   # Morado ODDS
SECONDARY = "#F1592A" # Naranja ODDS (opcional)


import io
import pandas as pd
from scipy import stats
def to_excel(code: str) -> io.BytesIO:
    if code not in store:
        return io.BytesIO()

    s = store[code]

    # A. Sesión estándar
    if s.get("tipo", "STD") == "STD":
        # Asegurar que todas las listas tienen la misma longitud
        n = min(
            len(s["ids"]),
            len(s["names"]),
            len(s["votes"]),
            len(s["comments"]),
            len(s.get("correos", []))  # solo si está presente
        )
        df = pd.DataFrame({
            "ID anónimo":    s["ids"][:n],
            "Nombre real":   s["names"][:n],
            "Correo":        s.get("correos", [""] * n)[:n],
            "Recomendación": [s["desc"]] * n,
            "Ronda":         [s["round"]] * n,
            "Voto":          s["votes"][:n],
            "Comentario":    s["comments"][:n],
            "Fecha":         [s["created_at"]] * n
        })

    elif s.get("tipo") == "GRADE_PKG":
        dominios = list(s["dominios"].keys())
        participantes = s["dominios"][dominios[0]]["names"]
        votos_por_dominio = {
            dom: s["dominios"][dom]["votes"]
            for dom in dominios
        }
        df = pd.DataFrame(votos_por_dominio, index=participantes).T
        df.index.name = "Dominio"
        df.columns.name = "Participante"

    buf = io.BytesIO()
    df.to_excel(buf, index=True)
    buf.seek(0)
    return buf


def crear_reporte_consolidado_recomendaciones(store: dict, history: dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import requests
    from io import BytesIO

    doc = Document()

    # — Logo en la cabecera —
    logo_url = (
        "https://static.wixstatic.com/media/89a9c2_ddc57311fc734357b9ea2b699e107ae2"
        "~mv2.png/v1/fill/w_90,h_54,al_c,q_85,usm_0.66_1.00_0.01/"
        "Logo%20versión%20principal.png"
    )
    resp = requests.get(logo_url)
    if resp.status_code == 200:
        img = BytesIO(resp.content)
        header_para = doc.sections[0].header.paragraphs[0]
        run = header_para.add_run()
        run.add_picture(img, width=Cm(4))
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # — Márgenes A4 —
    for sec in doc.sections:
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)

    # — Iterar cada sesión —    
    for code, s in store.items():
        votos = [v for v in s["votes"] if isinstance(v, (int, float))]
        total = len(votos)
        pct, med, lo, hi = consensus_pct(votos)*100, *median_ci(votos)
        quorum = s.get("n_participantes", 0)//2 + 1

        # Título
        h = doc.add_heading(level=1)
        h.add_run(f"Recomendación {code}").bold = True

        # Descripción y metadatos
        doc.add_paragraph(f"Descripción: {s['desc']}")
        doc.add_paragraph(f"Ronda: {s['round']}    Fecha: {s['created_at']}")

        # Tabla de métricas
        tbl = doc.add_table(rows=2, cols=4, style="Table Grid")
        hdr = tbl.rows[0].cells
        for i, title in enumerate(["Total votos", "% Consenso", "Mediana", "IC95%"]):
            p = hdr[i].paragraphs[0]
            run = p.add_run(title); run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row = tbl.rows[1].cells
        for i, val in enumerate([total, f"{pct:.1f}%", f"{med:.1f}", f"[{lo:.1f}, {hi:.1f}]"]):
            row[i].text = str(val)
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Estado de consenso
        if total < quorum:
            estado = "⚠️ Quórum no alcanzado"
        elif pct >= 80 and lo >= 7:
            estado = "✅ Consenso ALCANZADO"
        else:
            estado = "❌ No alcanzó consenso"

        p = doc.add_paragraph()
        p.add_run("Estado de consenso: ").bold = True
        p.add_run(estado)

        doc.add_page_break()

    # — Guardar en buffer y retornar —
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



def create_report(code: str) -> str:
    """
    Genera un reporte de texto plano con métricas y comentarios de la sesión actual
    (incluye también el historial de rondas anteriores si lo hay).
    """
    if code not in store:
        return "Sesión inválida"
    s = store[code]
    pct = consensus_pct(s["votes"]) * 100
    med, lo, hi = median_ci(s["votes"])
    # Cabecera
    lines = [
        f"REPORTE DE CONSENSO - Sesión {code}",
        f"Fecha de generación: {datetime.datetime.now():%Y-%m-%d %H:%M:%S}",
        "",
        f"Recomendación: {s['desc']}",
        f"Ronda actual: {s['round']}",
        f"Votos totales: {len(s['votes'])}",
        f"% Consenso: {pct:.1f}%",
        f"Mediana (IC95%): {med:.1f} [{lo:.1f}, {hi:.1f}]",
        "",
        "Comentarios:",
    ]
    # Comentarios de la ronda actual
    for pid, name, vote, com in zip(s["ids"], s["names"], s["votes"], s["comments"]):
        if com:
            lines.append(f"- {name} (ID {pid}): “{com}”")
    # Historial de rondas anteriores
    if code in history and history[code]:
        lines.append("\nHistorial de rondas anteriores:")
        for past in history[code]:
            ppct = consensus_pct(past["votes"]) * 100
            lines.append(
                f"  * Ronda {past['round']} [{past['created_at']}]: "
                f"%Consenso={ppct:.1f}%, Mediana={median_ci(past['votes'])[0]:.1f}"
            )
    return "\n".join(lines)


# Crear carpeta para guardar datos si no existe
DATA_DIR = "registro_data"
os.makedirs(DATA_DIR, exist_ok=True)

# Función: cargar registros desde CSV
def cargar_registros(nombre):
    path = os.path.join(DATA_DIR, f"{nombre}.csv")
    if os.path.exists(path):
        return pd.read_csv(path).to_dict("records")
    return []

# Función: guardar registros en CSV
def guardar_registros(nombre, registros):
    df = pd.DataFrame(registros)
    df.to_csv(os.path.join(DATA_DIR, f"{nombre}.csv"), index=False)

# Inicializar en session_state
if "registro_conflicto" not in st.session_state:
    st.session_state["registro_conflicto"] = cargar_registros("registro_conflicto")

if "registro_confidencialidad" not in st.session_state:
    st.session_state["registro_confidencialidad"] = cargar_registros("registro_confidencialidad")

# Lógica si la URL tiene ?registro=...
params = st.query_params
if "registro" in params:
    tipo = params.get("registro")
    st.set_page_config(page_title="Registro de Expertos", layout="centered")

    if tipo == "conflicto":
        st.title("🔐 Registro: Declaración de Conflictos de Interés")
        with st.form("form_conflicto_externo"):
            nombre = st.text_input("Nombre completo")
            institucion = st.text_input("Institución o afiliación")
            cargo = st.text_input("Cargo profesional")
            participa_en = st.multiselect("¿Participa actualmente en alguno de los siguientes?", [
                "Industria farmacéutica", "Investigación patrocinada", "Consultoría médica", "Autoría de guías clínicas", "Otro", "Ninguno"])
            tiene_conflicto = st.radio("¿Tiene un posible conflicto que pueda influir en esta recomendación?", ["No", "Sí"])
            detalle_conflicto = st.text_area("Describa brevemente su conflicto") if tiene_conflicto == "Sí" else ""
            confirma = st.checkbox("Declaro que la información es verídica y completa", value=False)
            submit = st.form_submit_button("Enviar")

            if submit:
                if not nombre or not confirma:
                    st.warning("Debe completar todos los campos obligatorios y aceptar la declaración.")
                else:
                    nuevo = {
                        "id": str(uuid.uuid4())[:8],
                        "nombre": nombre,
                        "institucion": institucion,
                        "cargo": cargo,
                        "participa_en": "; ".join(participa_en),
                        "conflicto": tiene_conflicto,
                        "detalle": detalle_conflicto,
                        "fecha": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    st.session_state["registro_conflicto"].append(nuevo)
                    guardar_registros("registro_conflicto", st.session_state["registro_conflicto"])
                    st.success("✅ Registro enviado correctamente. Puede cerrar esta ventana.")
        st.stop()

    elif tipo == "confidencialidad":
        st.title("📄 Registro: Acuerdo de Confidencialidad")
        with st.form("form_confidencialidad_externo"):
            nombre = st.text_input("Nombre completo")
            acepta1 = st.checkbox("Me comprometo a mantener la confidencialidad del contenido discutido y votado.")
            acepta2 = st.checkbox("Entiendo que no tengo derechos de autor sobre los productos resultantes del consenso.")
            submit = st.form_submit_button("Aceptar y registrar")

            if submit:
                if not nombre or not (acepta1 and acepta2):
                    st.warning("Debe completar el formulario y aceptar todas las condiciones.")
                else:
                    nuevo = {
                        "id": str(uuid.uuid4())[:8],
                        "nombre": nombre,
                        "fecha": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "acepta": True
                    }
                    st.session_state["registro_confidencialidad"].append(nuevo)
                    guardar_registros("registro_confidencialidad", st.session_state["registro_confidencialidad"])
                    st.success("✅ Registro enviado correctamente. Puede cerrar esta ventana.")
        st.stop()




# 2) Almacenamiento persistente utilizando session_state
# Esto asegura que las sesiones persistan incluso cuando Streamlit se reinicia
# Diccionario compartido en todo el servidor
@st.cache_resource
def get_store():
    return {}

store = get_store()
# Historia en memoria:
history = {}


# 3) Utilidades
def make_session(desc: str, scale: str) -> str:
    code = uuid.uuid4().hex[:6].upper()
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    store[code] = {
        "desc": desc,
        "scale": scale,
        "votes": [],
        "comments": [],
        "ids": [],
        "names": [],
        "created_at": timestamp,
        "round": 1
    }
    history[code] = []  # inicializamos el historial
    return code


def hash_id(name: str) -> str:
    return hashlib.sha256(name.encode()).hexdigest()[:8]

# Función para validar si un correo está autorizado para votar en una sesión privada
def correo_autorizado(correo: str, code: str) -> bool:
    if code in store:
        sesion = store[code]
        if sesion.get("privado", False):
            lista = sesion.get("correos_autorizados", [])
            return correo.lower().strip() in [c.lower().strip() for c in lista]
    return True

# Función para validar si un correo está autorizado para votar en una sesión privada
def correo_autorizado(correo: str, code: str) -> bool:
    if code in store:
        sesion = store[code]
        if sesion.get("privado", False):
            lista = sesion.get("correos_autorizados", [])
            return correo and correo.lower().strip() in [c.lower().strip() for c in lista]
    return True  # Si la sesión no es privada, siempre es autorizado

# Función para registrar el voto
def record_vote(code: str, vote, comment: str, name: str, correo: str = None):
    if code not in store:
        return None

    if not correo_autorizado(correo, code):
        return None

    s = store[code]
    pid = hashlib.sha256(name.encode()).hexdigest()[:8]

    if name and name in s["names"]:
        idx = s["names"].index(name)
        s["votes"][idx] = vote
        s["comments"][idx] = comment
        if "correos" in s and idx < len(s["correos"]):
            s["correos"][idx] = correo  # 🟢 actualiza el correo si ya existía
        return pid

    s["votes"].append(vote)
    s["comments"].append(comment)
    s["ids"].append(pid)
    s["names"].append(name)
    s.setdefault("correos", []).append(correo)  # 🟢 añade el correo nuevo
    return pid


def consensus_pct(votes):
    int_votes = [v for v in votes if isinstance(v, (int, float))]
    if not int_votes:
        return 0.0
    return sum(1 for v in int_votes if v >= 7) / len(int_votes)

def median_ci(votes):
    arr = np.array([v for v in votes if isinstance(v, (int, float))], dtype=float)
    n = arr.size
    if n == 0:
        return 0.0, 0.0, 0.0
    med = np.median(arr)
    if n < 2:
        return med, med, med
    try:
        res = stats.bootstrap((arr,), np.median,
                              confidence_level=0.95,
                              n_resamples=1000,
                              method="basic")
        lo, hi = res.confidence_interval
        if np.isnan(lo) or np.isnan(hi):
            lo = hi = med
    except:
        lo = hi = med
    return med, lo, hi

def get_base_url():
    # URL específica para aplicación en Streamlit Cloud
    return "https://consenso-expertos-sfpqj688ihbl7m6tgrdmwb.streamlit.app"

def create_qr_code_url(code: str):
    base_url = get_base_url()
    # Elimina slashes finales para evitar doble slash
    base_url = base_url.rstrip('/')
    # Construye URL correctamente
    return f"{base_url}/?session={code}"

def make_qr(code: str) -> io.BytesIO:
    url = create_qr_code_url(code)
    
    buf = io.BytesIO()
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,  # Nivel más alto de corrección de errores
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf

def get_qr_code_image_html(code):
    buf = make_qr(code)
    img_str = base64.b64encode(buf.getvalue()).decode("utf-8")
    url = create_qr_code_url(code)
    html = f"""
    <div style="text-align: center; margin-bottom: 20px;">
        <img src="data:image/png;base64,{img_str}" width="200">
        <p style="margin-top: 10px; font-size: 0.8rem;">URL: <a href="{url}" target="_blank">{url}</a></p>
    </div>
    """
    return html



def crear_reporte_consolidado_recomendaciones(store: dict, history: dict) -> io.BytesIO:
    """
    Genera un .docx con, para cada recomendación:
      - Logo alineado a la derecha en la cabecera
      - Encabezado con el código
      - Descripción
      - Fecha de creación
      - Tabla de métricas (Total votos, % Consenso, Mediana, IC95%)
      - Estado de consenso
    """
    doc = Document()



# ——————————————————————————————
#  Integración en Streamlit
# ——————————————————————————————
def integrar_reporte_todas_recomendaciones():
    st.subheader(" Descargar Reporte Consolidado de Recomendaciones")

    if not store:
        st.info("No hay recomendaciones registradas aún.")
        return

    if st.button("⬇️ Generar y Descargar .docx"):
        buf = crear_reporte_consolidado_recomendaciones(store, history)
        nombre = f"reporte_recomendaciones_{datetime.datetime.now():%Y%m%d}.docx"
        st.download_button(
            label="Descargar Documento",
            data=buf,
            file_name=nombre,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )



# ─────────────────────────────────────────────────────────────
# 5)  Página de votación (se adapta al tipo de sesión)
# ─────────────────────────────────────────────────────────────
# 5) Página de votación (adaptable al tipo de sesión)
# ─────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────
# 5) Página de votación (se adapta al tipo de sesión)
# ─────────────────────────────────────────────────────────────
# ——————————————————————————————
# Manejo de la página de votación según ?session=…
# ——————————————————————————————
import streamlit as st
import hashlib
# … tus otros imports …

# ——————————————————————————————
# Pantalla de votación (oculta el panel de administración)
# ——————————————————————————————
# ——————————————————————————————
# Manejo de la página de votación según ?session=…
# ——————————————————————————————
params = st.query_params

if "session" in params:
    # Extraer y limpiar el código de sesión
    raw = params.get("session")
    code = raw[0] if isinstance(raw, list) else raw
    code = str(code).strip().upper()

    odds_header()

    # Ocultar barra lateral en la página de votación
    st.markdown("""
        <style>
          [data-testid="stSidebar"] { display: none !important; }
          [data-testid="collapsedControl"] { display: none !important; }
        </style>
    """, unsafe_allow_html=True)

    s = store.get(code)
    if not s:
        st.error(f"Sesión inválida: {code}")
        st.stop()

    tipo = s.get("tipo", "STD")
    es_privada = s.get("privado", False)

    st.subheader(f"Panel de votación — Sesión {code}")

    # Capturar nombre
    name = st.text_input("Nombre completo (nombre y apellido) del participante:")

    # Capturar correo solo si la sesión es privada
    correo = None
    if es_privada:
        correo = st.text_input("Correo electrónico (obligatorio):")
        if not name or not correo:
            st.warning("Ingrese nombre y correo electrónico para continuar.")
            st.stop()
        if not correo_autorizado(correo, code):
            st.error("❌ El correo ingresado no está autorizado para participar en esta sesión privada.")
            st.stop()
    else:
        if not name:
            st.warning("Ingrese su nombre para continuar.")
            st.stop()

    # Verificar si ya votó
    ya_voto = (
        (tipo == "STD" and name in s["names"]) or
        (tipo == "GRADE_PKG" and name in s["dominios"]["prioridad_problema"]["names"])
    )
    if ya_voto:
        st.success("✅ Ya registró su participación.")
        st.stop()

    # Tipo estándar: Voto por paquete de recomendaciones
    if tipo == "STD":
        st.markdown("""
        <div style="margin-top: 10px; padding: 10px; background-color: #f0f2f6; border-left: 4px solid #662D91; border-radius: 5px;">
        ⚠️ <strong>Importante:</strong> Lea todas las recomendaciones antes de emitir su voto.<br>
        Al finalizar el recorrido podrá emitir un único voto global para todo el paquete.
        </div>
        """, unsafe_allow_html=True)

        import re
        def separar_recomendaciones(texto):
            partes = re.split(r'\s*\d+\.\s*', str(texto))
            return [p.strip() for p in partes if p.strip()]

        if "lista_recos" not in st.session_state:
            st.session_state.lista_recos = separar_recomendaciones(s["desc"])
            st.session_state.reco_index = 0

        index = st.session_state.reco_index
        total = len(st.session_state.lista_recos)
        reco_actual = st.session_state.lista_recos[index]

        if "imagenes_relacionadas" in s and s["imagenes_relacionadas"]:
            st.markdown("Haga click sobre las lupas si quiere ver las tablas relacionadas con esta/s recomendacion/es")
            for i, img_bytes in enumerate(s["imagenes_relacionadas"]):
                with st.expander(f"🔍 Ver tablas {i+1}"):
                    st.image(img_bytes, use_container_width=True)

        col_left, col_center, col_right = st.columns([1, 8, 1])
        with col_left:
            if st.button("⬅️", key="anterior", disabled=(index == 0)):
                st.session_state.reco_index -= 1
                st.rerun()
        with col_center:
            st.markdown(f"**Recomendación {index+1} de {total}**")
            st.markdown(reco_actual)
        with col_right:
            if st.button("➡️", key="siguiente", disabled=(index == total - 1)):
                st.session_state.reco_index += 1
                st.rerun()

        # Solo permitir votar en la última recomendación
        if index == total - 1:
            st.markdown("---")
            st.markdown("**1–3 Desacuerdo • 4–6 Neutral • 7–9 Acuerdo**")
            voto = st.slider("Su voto global para todas las recomendaciones:", 1, 9, 5, key="voto_final")
            comentario = st.text_area("Comentario (opcional):", key="comentario_final")

            if st.button("✅ Enviar voto"):
                pid = hashlib.sha256(name.encode()).hexdigest()[:8]

                if name not in s["names"]:
                    s["names"].append(name)
                    s["ids"].append(pid)

                s["votes"].append(voto)
                s["comments"].append(comentario)
                s.setdefault("correos", []).append(correo)

                st.balloons()
                st.success(f"🎉 Su voto ha sido registrado. ID: `{pid}`")

                for k in ["lista_recos", "reco_index"]:
                    st.session_state.pop(k, None)

                st.stop()

    # Detener cualquier render posterior innecesario
    st.stop()


# … aquí continúa el resto de tu aplicación (panel de administración, sidebar, etc.) …

# 6) Panel de administración
odds_header()
# Logo en la barra lateral
logo_url = "https://static.wixstatic.com/media/89a9c2_ddc57311fc734357b9ea2b699e107ae2~mv2.png/v1/fill/w_90,h_54,al_c,q_85,usm_0.66_1.00_0.01/Logo%20versi%C3%B3n%20principal.png"
st.sidebar.image(logo_url, width=80)

st.sidebar.title("Panel de Control")
st.sidebar.markdown("### ODDS Epidemiology")
menu = st.sidebar.selectbox("Navegación", ["Inicio", "Crear Recomendación", "Dashboard", "Crear Paquete GRADE", "Reporte Consolidado"])

if menu == "Inicio":
    st.markdown("## Bienvenido al Sistema de votación para Consenso de expertos de ODDS Epidemiology")
    
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("""
    
    
    Utilice el panel de navegación para comenzar.
    """)
 
    
    

# ──────────────────────────────────────────────────────────────────────────────
#  BLOQUE DEL PANEL: "Crear Recomendación"
# ──────────────────────────────────────────────────────────────────────────────
elif menu == "Crear Recomendación":
    st.subheader("Crear Nueva Recomendación")
    st.markdown('<div class="card">', unsafe_allow_html=True)

    import re
    def separar_recomendaciones(texto):
        partes = re.split(r'\s*\d+\.\s*', str(texto))
        return [p.strip() for p in partes if p.strip()]

    st.markdown("### Cargar recomendaciones desde Excel")
    if "uploader_key" not in st.session_state:
        st.session_state.uploader_key = 0

    excel_file = st.file_uploader(
        "Suba archivo .xlsx/.xls con columnas 'ronda' y 'recomendacion'",
        type=["xlsx", "xls"],
        key=f"excel_{st.session_state.uploader_key}"
    )

    if excel_file and "recomendaciones_precargadas" not in st.session_state:
        try:
            df = pd.read_excel(excel_file, engine="openpyxl")
            df.columns = df.columns.str.strip().str.lower()
            if "recomendacion" not in df.columns:
                st.error("El archivo debe tener al menos una columna llamada 'recomendacion'.")
            else:
                preguntas = df["recomendacion"].dropna().tolist()
                modo = st.radio("¿Cómo desea proceder con las recomendaciones?", ["Usar todas las recomendaciones", "Seleccionar recomendaciones manualmente"])
                if modo == "Usar todas las recomendaciones":
                    texto_final = "\n".join([f"{i+1}. {rec}" for i, rec in enumerate(preguntas)])
                    st.session_state["ronda_precargada"] = df["ronda"].iloc[0] if "ronda" in df.columns else ""
                    st.session_state["recomendaciones_precargadas"] = texto_final.strip()
                    st.success(f"✅ {len(preguntas)} recomendaciones cargadas para la sesión.")
                elif modo == "Seleccionar recomendaciones manualmente":
                    seleccionadas = st.multiselect("Seleccione las recomendaciones que desea incluir:", options=preguntas)
                    if seleccionadas:
                        texto_final = "\n".join([f"{i+1}. {rec}" for i, rec in enumerate(seleccionadas)])
                        st.session_state["ronda_precargada"] = df["ronda"].iloc[0] if "ronda" in df.columns else ""
                        st.session_state["recomendaciones_precargadas"] = texto_final.strip()
                        st.success(f"✅ {len(seleccionadas)} recomendaciones seleccionadas para la sesión.")
        except Exception as e:
            st.error(f"Error al leer el archivo: {e}")

    if "recomendaciones_precargadas" in st.session_state and st.button("❌ Quitar archivo cargado"):
        for k in ["ronda_precargada", "recomendaciones_precargadas"]:
            st.session_state.pop(k, None)
        st.session_state.uploader_key += 1
        st.experimental_rerun()

    st.markdown("<hr>", unsafe_allow_html=True)

    # —— Formulario de creación manual ——
    with st.form("create_form", clear_on_submit=True):
        codigo_manual = st.text_input("Código personalizado para la sesión (opcional):").strip().upper()
        nombre_ronda = st.text_input("Nombre de la ronda:", value=st.session_state.pop("ronda_precargada", ""))
        desc = st.text_area("Recomendaciones a evaluar:", value=st.session_state.pop("recomendaciones_precargadas", ""), height=300)
        scale = st.selectbox("Escala de votación:", ["Likert 1-9", "Sí/No"])
        n_participantes = st.number_input("¿Cuántos participantes están habilitados para votar?", min_value=1, step=1)
        es_privada = st.checkbox("¿Esta recomendación será privada?")
        imagenes_subidas = st.file_uploader("📷 Cargar imágenes relacionadas (opcional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

        correos_autorizados = []
        archivo_correos = st.file_uploader("📧 Lista de correos autorizados (CSV con columna 'correo')", type=["csv"])
        if archivo_correos:
            try:
                df_correos = pd.read_csv(archivo_correos)
                if "correo" in df_correos.columns:
                    correos_autorizados = df_correos["correo"].astype(str).str.strip().tolist()
                    st.success(f"{len(correos_autorizados)} correos cargados.")
                else:
                    st.error("El CSV debe contener una columna llamada 'correo'.")
            except Exception as e:
                st.error(f"No se pudo leer el CSV: {e}")

        st.markdown("""
        <div class="helper-text">
        Escala Likert 1‑9:<br>
        • 1‑3 Desacuerdo • 4‑6 Neutral • 7‑9 Acuerdo<br>
        Se alcanza consenso cuando ≥80 % de votos son ≥7 y hay quórum (mitad + 1).
        </div>
        """, unsafe_allow_html=True)

        if st.form_submit_button("Crear Recomendación"):
            if not desc:
                st.warning("Por favor, ingrese la recomendación.")
                st.stop()

            code = codigo_manual if codigo_manual else uuid.uuid4().hex[:6].upper()
            if code in store:
                st.error("❌ Ese código ya está en uso. Elija otro.")
                st.stop()

            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            descripcion_final = f"{desc} ({nombre_ronda})" if nombre_ronda else desc

            store[code] = {
                "desc": descripcion_final,
                "scale": scale,
                "votes": [],
                "comments": [],
                "ids": [],
                "names": [],
                "created_at": timestamp,
                "round": 1,
                "is_active": True,
                "n_participantes": int(n_participantes),
                "privado": es_privada,
                "correos_autorizados": correos_autorizados,
                "imagenes_relacionadas": [img.getvalue() for img in imagenes_subidas] if imagenes_subidas else []
            }
            history[code] = []

            st.success("✅ Sesión creada exitosamente.")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"""
                <div class="metric-card">
                  <div class="metric-label">Código de sesión</div>
                  <div class="metric-value">{code}</div>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(get_qr_code_image_html(code), unsafe_allow_html=True)

            url = create_qr_code_url(code)
            st.info(f"URL para compartir: {url}")
            st.write(f"[Abrir página de votación]({url})")
            st.markdown("</div>", unsafe_allow_html=True)



elif menu == "Dashboard":
    st.subheader("Dashboard en Tiempo Real")
    st_autorefresh(interval=5000, key="refresh_dashboard")

    # Selección de sesión
    active_sessions = [k for k, v in store.items() if v.get("is_active", True)]
    if not active_sessions:
        st.info("No hay sesiones activas.")
        st.stop()
    code = st.selectbox("Seleccionar sesión activa:", active_sessions)
    if not code:
        st.stop()

    # Datos de la sesión
    s = store.get(code)
    if not s:
        st.error("Código de sesión no encontrado.")
        st.stop()

    # Importante: solo tomar un voto por participante
    votes = [
        s["votes"][i] for i in range(len(s["names"]))
        if i < len(s["votes"]) and isinstance(s["votes"][i], (int, float))
    ]

    n = len(votes)
    media = np.mean(votes) if n > 0 else 0.0
    desv_std = np.std(votes, ddof=1) if n > 1 else 0.0
    mediana, lo, hi = median_ci(votes)
    pct = consensus_pct(votes) * 100
    quorum = s.get("n_participantes", 0) // 2 + 1
    votos_actuales = len(set(s["names"]))

    # Tres columnas: Resumen | Métricas | Gráfico
    col_res, col_kpi, col_chart = st.columns([2, 1, 3])

    # --- Columna 1: Resumen ---
    with col_res:
        if st.button("Finalizar esta sesión"):
            store[code]["is_active"] = False
            history.setdefault(code, []).append(copy.deepcopy(s))
            st.success("✅ Sesión finalizada.")
            st.rerun()
        st.markdown(f"""
        **Recomendación:** {s['desc']}  
        **Ronda actual:** {s['round']}  
        **Creada:** {s['created_at']}  
        **Votos esperados:** {s.get('n_participantes','?')}  
        **Quórum:** {quorum}  
        **Votos recibidos:** {votos_actuales}
        """)

    # --- Columna 2: Métricas ---
    with col_kpi:
        st.markdown(card_html("Media", f"{media:.2f}"), unsafe_allow_html=True)
        st.markdown(card_html("Desv. estándar", f"{desv_std:.2f}"), unsafe_allow_html=True)
        st.markdown(card_html("% Consenso", f"{pct:.1f}%"), unsafe_allow_html=True)
        if n > 0:
            st.markdown(card_html("Mediana (IC95%)", f"{mediana:.1f} [{lo:.1f}, {hi:.1f}]"), unsafe_allow_html=True)

    # --- Columna 3: Histograma ---
    with col_chart:
        if votos_actuales:
            df = pd.DataFrame({"Voto": votes})
            fig = px.histogram(
                df, x="Voto", nbins=9,
                labels={"Voto": "Escala 1–9", "count": "Frecuencia"},
                color_discrete_sequence=[PRIMARY]
            )
            fig.update_traces(marker_line_width=0)
            fig.update_layout(
                bargap=0.4,
                xaxis=dict(tickmode="linear", tick0=1, dtick=1),
                margin=dict(t=30, b=0, l=0, r=0),
                height=300,
                plot_bgcolor="rgba(0,0,0,0)",
                paper_bgcolor="rgba(0,0,0,0)"
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("🔍 Aún no hay votos para mostrar.")

    # Estado de consenso
    st.markdown("---")
    if votos_actuales < quorum:
        st.info(f"🕒 Quórum no alcanzado ({votos_actuales}/{quorum})")
    else:
        if pct >= 80 and 7 <= mediana <= 9 and 7 <= lo <= 9 and 7 <= hi <= 9:
            st.success("✅ CONSENSO ALCANZADO (mediana + IC95%)")
        elif pct >= 80:
            st.success("✅ CONSENSO ALCANZADO (% votos)")
        elif pct <= 20 and 1 <= mediana <= 3 and 1 <= lo <= 3 and 1 <= hi <= 3:
            st.error("❌ NO APROBADO (mediana + IC95%)")
        elif sum(1 for v in votes if v <= 3) >= 0.8 * votos_actuales:
            st.error("❌ NO APROBADO (% votos)")
        else:
            st.warning("⚠️ NO SE ALCANZÓ CONSENSO")

    # Acciones de exportación
    st.subheader("Acciones y Exportación")
    if st.button("Iniciar nueva ronda"):
        history.setdefault(code, []).append(copy.deepcopy(s))
        st.session_state.modify_recommendation = True
        st.session_state.current_code = code

    c1, c2 = st.columns(2)
    with c1:
        buf = to_excel(code)
        if buf:
            st.download_button("⬇️ Descargar Excel", data=buf.getvalue(),
                               file_name=f"consenso_{code}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.warning("⚠️ No hay datos disponibles para exportar en Excel.")

    with c2:
        st.download_button("⬇️ Descargar TXT", create_report(code),
                           file_name=f"reporte_{code}.txt")

    # Comentarios
    if s.get("comments"):
        st.subheader("Comentarios de Participantes")
        for pid, name, vote, com in zip(s["ids"], s["names"], s["votes"], s["comments"]):
            if com:
                st.markdown(f"**{name}** (ID:{pid}) — Voto: {vote}\n> {com}")


elif menu == "Crear Paquete GRADE":
    st.subheader("Crear / Descargar Paquetes GRADE")

    # ——— Crear nuevo paquete ———
    st.markdown("#### 1. Crear nuevo paquete")
    opciones = list(store.keys())
    sel = st.multiselect(
        "Elige las recomendaciones para el paquete:",
        opciones,
        format_func=lambda c: f"{c} – {store[c]['desc']}"
    )
    n_part = st.number_input("¿Cuántos expertos?", min_value=1, step=1)
    if st.button("Crear Paquete"):
        code = uuid.uuid4().hex[:6].upper()
        dominios = {
            dom: {"ids":[], "names":[], "votes":[], "comments":[], "opciones": DOMINIOS_GRADE[dom]}
            for dom in DOMINIOS_GRADE
        }
        store[code] = {
            "tipo": "GRADE_PKG",
            "desc": f"Paquete de {len(sel)} recomendaciones",
            "recs": sel,
            "dominios": dominios,
            "n_participantes": n_part,
            "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "is_active": True
        }
        history[code] = []
        st.success(f"Paquete GRADE creado con código **{code}**")
        st.markdown(get_qr_code_image_html(code), unsafe_allow_html=True)
        st.info("🔗 Comparte este QR para que los expertos voten.")

    st.markdown("---")

    # ——— Descargar resultados de paquetes existentes ———
    st.markdown("#### 2. Descargar resultados de paquetes existentes")
    # Filtramos sólo los que ya tienen al menos un voto (para no listar paquetes vacíos)
    paquetes = [
        c for c, s in store.items()
        if s.get("tipo") == "GRADE_PKG"
        and len(next(iter(s["dominios"].values()))["votes"]) > 0
    ]

    if paquetes:
        sel_pkg = st.selectbox(
            "Selecciona un paquete para descargar:",
            paquetes,
            format_func=lambda c: f"{c} – {len(store[c]['dominios']['prioridad_problema']['votes'])} votos"
        )
        buf2 = to_excel(sel_pkg)
        st.download_button(
            "⬇️ Descargar Excel del paquete",
            data=buf2,
            file_name=f"GRADE_{sel_pkg}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.info("No hay paquetes con votos para descargar.")


elif menu == "Reporte Consolidado":
    st.header("📊 Reporte Consolidado")
    st.subheader("Libro Excel (.xlsx)")

    # 1. Generar el buffer con todas las hojas
    buf_xls = crear_excel_consolidado(store, history)

    # 2. Debug: comprobar que realmente tiene la hoja “Métricas”
    import pandas as pd
    xls = pd.ExcelFile(buf_xls)  
    st.write("📑 Hojas en el Excel:", xls.sheet_names)

    # 3. Botón de descarga
    st.download_button(
        label="⬇️ Descargar Reporte Consolidado (.xlsx)",
        data=buf_xls.getvalue(),
        file_name=f"reporte_consolidado_{datetime.datetime.now():%Y%m%d}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# Cargar estado
state_upload = st.sidebar.file_uploader("Cargar Estado", type=["txt"])
if state_upload is not None:
    try:
        content = state_upload.read().decode()
        decoded = base64.b64decode(content).decode()
        import ast
        state_data = ast.literal_eval(decoded)

        if "sessions" in state_data and "history" in state_data:
            store.clear()
            store.update(state_data["sessions"])
            history.clear()
            history.update(state_data["history"])
            st.sidebar.success("Estado restaurado correctamente.")
            st.rerun()
        else:
            st.sidebar.error("El archivo no contiene datos válidos.")
    except Exception as e:
        st.sidebar.error(f"Error al cargar el estado: {str(e)}")

elif menu == "Registro Previo":
    st.title("Registro Previo - Panel de Consenso")
    st.markdown("Comparta los siguientes enlaces con los participantes para que completen sus registros antes de iniciar el consenso.")

    #  Conflictos de Interés
    st.markdown("###  Declaración de Conflictos de Interés")
    url_conflicto = "https://consenso-expertos-sfpqj688ihbl7m6tgrdmwb.streamlit.app/?registro=conflicto"
    st.code(url_conflicto)
    qr_conflicto = qrcode.make(url_conflicto)
    buf1 = io.BytesIO()
    qr_conflicto.save(buf1, format="PNG")
    img1 = base64.b64encode(buf1.getvalue()).decode()
    st.markdown(f'<img src="data:image/png;base64,{img1}" width="180">', unsafe_allow_html=True)

    # 📄 Confidencialidad
    st.markdown("---")
    st.markdown("###  Compromiso de Confidencialidad")
    url_confid = "https://consenso-expertos-sfpqj688ihbl7m6tgrdmwb.streamlit.app/?registro=confidencialidad"
    st.code(url_confid)
    qr_confid = qrcode.make(url_confid)
    buf2 = io.BytesIO()
    qr_confid.save(buf2, format="PNG")
    img2 = base64.b64encode(buf2.getvalue()).decode()
    st.markdown(f'<img src="data:image/png;base64,{img2}" width="180">', unsafe_allow_html=True)

    # 📥 Exportar datos recibidos
    st.markdown("---")
    st.subheader(" Exportar registros recibidos")

    col1, col2 = st.columns(2)

    if st.session_state["registro_conflicto"]:
        df1 = pd.DataFrame(st.session_state["registro_conflicto"])
        with col1:
            st.download_button("⬇️ Descargar Conflictos", df1.to_csv(index=False).encode(), file_name="conflictos.csv")
    else:
        with col1:
            st.info("Sin registros aún.")

    if st.session_state["registro_confidencialidad"]:
        df2 = pd.DataFrame(st.session_state["registro_confidencialidad"])
        with col2:
            st.download_button("⬇️ Descargar Confidencialidad", df2.to_csv(index=False).encode(), file_name="confidencialidad.csv")
    else:
        with col2:
            st.info("Sin registros aún.")

    # 🗑️ Borrar registros (SOLO dentro del menú Registro Previo)
    st.markdown("---")
    st.subheader("🗑️ Borrar registros")

    if st.button("❌ Borrar todos los registros de conflicto y confidencialidad"):
        st.session_state["registro_conflicto"] = []
        st.session_state["registro_confidencialidad"] = []

        try:
            os.remove(os.path.join(DATA_DIR, "registro_conflicto.csv"))
            os.remove(os.path.join(DATA_DIR, "registro_confidencialidad.csv"))
        except FileNotFoundError:
            pass

        st.success("Registros eliminados correctamente.")




# Créditos
st.sidebar.markdown("---")
st.sidebar.markdown("**ODDS Epidemiology**")
st.sidebar.markdown("v1.0.0 - 2025")
st.sidebar.markdown("© Todos los derechos reservados")
